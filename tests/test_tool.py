"""Pytest suite: search recency, fuzzy matching, template population,
resumable indexing. No API calls — extraction is stubbed."""

import sqlite3
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from common import SCHEMA
import make_quote
import index_archive


@pytest.fixture
def db(tmp_path):
    """Temp DB seeded with VLV-2043 in three quotes across three years."""
    path = tmp_path / "test.db"
    con = sqlite3.connect(path)
    con.execute("PRAGMA foreign_keys = ON")
    con.executescript(SCHEMA)
    rows = [
        ("a.docx", "Q-2022-1", "2022-03-01", "Old Co"),
        ("b.docx", "Q-2024-1", "2024-06-15", "Mid Co"),
        ("c.docx", "Q-2026-1", "2026-01-20", "New Co"),
    ]
    items = [
        (1, "VLV-2043", "2\" Stainless Ball Valve", "Old desc.", 184.50, 2),
        (2, "VLV-2043", "2\" Stainless Ball Valve", "Mid desc.", 205.10, 1),
        (3, "VLV-2043", "2\" Stainless Ball Valve", "New desc.", 231.99, 4),
        (2, "PMP-118", "Centrifugal Transfer Pump 1.5HP", "Pump desc.", 1310.00, 1),
    ]
    con.executemany(
        "INSERT INTO quotes (file, quote_number, quote_date, customer) "
        "VALUES (?,?,?,?)", rows)
    con.executemany(
        "INSERT INTO line_items (quote_id, part_number, product_name, "
        "description, unit_price, quantity) VALUES (?,?,?,?,?,?)", items)
    con.commit()
    con.close()
    return path


def test_search_returns_most_recent(db):
    con = sqlite3.connect(db)
    hit = make_quote.lookup_part(con, "vlv-2043")  # case-insensitive too
    assert hit is not None
    assert hit["quote_number"] == "Q-2026-1"
    assert hit["unit_price"] == 231.99
    assert hit["description"] == "New desc."


def test_missing_description_backfills_from_older_quote(db):
    """Most recent price wins, but a null description/name is backfilled
    from the most recent occurrence that has one."""
    con = sqlite3.connect(db)
    con.execute("PRAGMA foreign_keys = ON")
    con.execute("INSERT INTO quotes (file, quote_number, quote_date, customer) "
                "VALUES ('d.docx', 'Q-2026-2', '2026-05-01', 'Newest Co')")
    qid = con.execute("SELECT id FROM quotes WHERE file='d.docx'").fetchone()[0]
    con.execute("INSERT INTO line_items (quote_id, part_number, product_name, "
                "description, unit_price, quantity) "
                "VALUES (?, 'VLV-2043', NULL, NULL, 240.00, 1)", (qid,))
    con.commit()

    hit = make_quote.lookup_part(con, "VLV-2043")
    assert hit["unit_price"] == 240.00           # newest quote's price
    assert hit["quote_number"] == "Q-2026-2"
    assert hit["description"] == "New desc."     # backfilled from 2026-01
    assert hit["product_name"] == "2\" Stainless Ball Valve"


def test_search_missing_part_returns_none(db):
    con = sqlite3.connect(db)
    assert make_quote.lookup_part(con, "XYZ-0000") is None


def test_fuzzy_matching_suggests_close_parts(db):
    con = sqlite3.connect(db)
    sugg = make_quote.suggest(con, "VLV-2034")   # transposed digits
    assert "VLV-2043" in sugg
    sugg2 = make_quote.suggest(con, "pmp118")    # missing hyphen, lowercase
    assert "PMP-118" in sugg2


def test_template_population(db, tmp_path, monkeypatch):
    out_dir = tmp_path / "out"
    path, found, missing = make_quote.generate(
        ["VLV-2043", "PMP-118"], "Jane Doe", "Apex Fabrication LLC",
        quantities=[3, 1], db_path=db, out_dir=out_dir)
    assert path is not None and path.exists()
    assert not missing

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
    all_text = text + "\n" + "\n".join(cells)

    assert "Jane Doe" in all_text
    assert "Apex Fabrication LLC" in all_text
    assert "VLV-2043" in cells
    assert "PMP-118" in cells
    assert "$231.99" in cells          # most recent VLV price
    assert "$1,310.00" in cells
    # 3 x 231.99 + 1 x 1310.00 = 2005.97
    assert "$2,005.97" in all_text
    # No unreplaced tokens anywhere.
    assert "{{" not in all_text


def test_generate_fails_gracefully_on_missing_part(db, tmp_path):
    path, found, missing = make_quote.generate(
        ["VLV-2043", "XYZ-9999"], "Jane", "Co",
        db_path=db, out_dir=tmp_path / "out")
    assert path is None                      # nothing written by default
    assert len(found) == 1
    assert missing[0][0] == "XYZ-9999"

    # --allow-partial writes with the found item only.
    path2, found2, _ = make_quote.generate(
        ["VLV-2043", "XYZ-9999"], "Jane", "Co",
        db_path=db, out_dir=tmp_path / "out2", allow_partial=True)
    assert path2 is not None and path2.exists()


def test_resumable_indexing(tmp_path):
    """Second run over the same files must not re-extract anything."""
    con = sqlite3.connect(tmp_path / "idx.db")
    con.execute("PRAGMA foreign_keys = ON")
    con.executescript(SCHEMA)

    # Two fake "documents" (plain docx files so extract_text works).
    files = []
    for name in ("q1.docx", "q2.docx"):
        p = tmp_path / name
        d = Document()
        d.add_paragraph("Quote " + name)
        d.save(str(p))
        files.append(p)

    calls = []

    def fake_extractor(text):
        calls.append(text)
        return ({"quote_number": "Q-1", "quote_date": "2025-01-01",
                 "customer": "C", "line_items": []}, 10, 5)

    n1, s1, f1, _, _ = index_archive.run_index(con, fake_extractor, files)
    assert (n1, s1, f1) == (2, 0, [])
    assert len(calls) == 2

    n2, s2, f2, _, _ = index_archive.run_index(con, fake_extractor, files)
    assert (n2, s2) == (0, 2)
    assert len(calls) == 2                   # no new extraction calls

    # force=True re-indexes everything.
    n3, _, _, _, _ = index_archive.run_index(con, fake_extractor, files, force=True)
    assert n3 == 2
    assert len(calls) == 4
