"""Shared helpers: env loading, SQLite connection/schema, document text extraction."""

import os
import sqlite3
from pathlib import Path

ROOT = Path(__file__).parent
DB_PATH = ROOT / "data" / "quotes.db"
ARCHIVE_DIR = ROOT / "archive"

SCHEMA = """
CREATE TABLE IF NOT EXISTS quotes (
    id            INTEGER PRIMARY KEY,
    file          TEXT NOT NULL UNIQUE,
    quote_number  TEXT,
    quote_date    TEXT,            -- ISO YYYY-MM-DD
    customer      TEXT,
    indexed_at    TEXT NOT NULL DEFAULT (datetime('now')),
    input_tokens  INTEGER DEFAULT 0,
    output_tokens INTEGER DEFAULT 0
);
CREATE TABLE IF NOT EXISTS line_items (
    id           INTEGER PRIMARY KEY,
    quote_id     INTEGER NOT NULL REFERENCES quotes(id) ON DELETE CASCADE,
    part_number  TEXT,
    product_name TEXT,
    description  TEXT,
    unit_price   REAL,
    quantity     INTEGER
);
CREATE INDEX IF NOT EXISTS idx_line_items_part ON line_items(part_number);
"""


def load_env():
    """Load KEY=VALUE pairs from .env next to this file (no external deps).

    Existing environment variables win over .env values.
    """
    env_file = ROOT / ".env"
    if not env_file.exists():
        return
    for line in env_file.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        k, v = k.strip(), v.strip().strip('"').strip("'")
        if k and k not in os.environ:
            os.environ[k] = v


def require_api_key():
    load_env()
    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        raise SystemExit(
            "ANTHROPIC_API_KEY not found.\n"
            "Set it one of these ways:\n"
            '  PowerShell (this session):  $env:ANTHROPIC_API_KEY = "sk-ant-..."\n'
            '  PowerShell (permanent):     [Environment]::SetEnvironmentVariable('
            '"ANTHROPIC_API_KEY", "sk-ant-...", "User")\n'
            "  Or create a .env file next to this script containing:\n"
            "      ANTHROPIC_API_KEY=sk-ant-...\n"
        )
    return key


def db_connect(path=DB_PATH):
    path.parent.mkdir(exist_ok=True)
    con = sqlite3.connect(path)
    con.execute("PRAGMA foreign_keys = ON")
    con.executescript(SCHEMA)
    return con


def extract_text(path: Path) -> str:
    """Plain-text dump of a .docx (python-docx) or .pdf (pdfplumber)."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        import pdfplumber
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n".join(pages)
    raise ValueError(f"Unsupported file type: {path}")
