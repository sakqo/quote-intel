"""Generate a fake archive of past sales quotes for Meridian Industrial Supply Co.

Creates N quotes (default 25, --count to scale) spanning 2022-2026, each saved
as BOTH .docx and .pdf, with deliberate messiness:
  - 3 layout styles: table-based, paragraph-based, letter-style
  - inconsistent field labels (Part No. / PN / Part # / Item Number ...)
  - some line items missing description or quantity; one quote missing customer
  - price formats vary ($1,204.50 vs 1204.5 USD vs USD 1,204.50 ...)
  - a handful of "popular" part numbers recur across years with price drift

Ground truth for every file/quote/line item is written to data/answer_key.json.
The extractor must NEVER read that file; it exists only for validation.
"""

import argparse
import json
import random
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
)

ROOT = Path(__file__).parent

# ---------------------------------------------------------------- catalog ---

CATALOG = [
    # part, name, description (2-4 sentences), base price (2022), annual drift
    ("VLV-2043", "2\" Stainless Ball Valve",
     "Full-port two-piece ball valve in 316 stainless steel with PTFE seats. "
     "Rated to 1000 PSI WOG and suitable for steam service up to 150 PSI. "
     "Locking lever handle included as standard.", 184.50, 0.06),
    ("VLV-1120", "3/4\" Bronze Gate Valve",
     "Class 125 bronze gate valve with threaded NPT ends and a non-rising stem. "
     "Solid wedge disc for positive shutoff in water and low-pressure steam lines. "
     "Meets MSS SP-80 specifications.", 42.75, 0.05),
    ("PMP-118", "Centrifugal Transfer Pump 1.5HP",
     "End-suction centrifugal pump with cast iron volute and bronze impeller. "
     "Delivers up to 85 GPM at 60 feet of head on a 1.5 HP TEFC motor. "
     "Mechanical seal is carbon-ceramic rated for continuous duty. "
     "Suitable for clear water transfer and light industrial coolant loops.", 1204.50, 0.07),
    ("PMP-3305", "Diaphragm Metering Pump",
     "Solenoid-driven diaphragm metering pump for chemical dosing applications. "
     "Adjustable output from 0.5 to 12 GPH with PVDF liquid end for aggressive media. "
     "Includes foot valve, injection fitting, and 6 feet of suction tubing.", 689.00, 0.04),
    ("BRG-7742", "Spherical Roller Bearing 65mm",
     "Double-row spherical roller bearing with a 65 mm bore and brass cage. "
     "Self-aligning design tolerates up to 2 degrees of shaft misalignment. "
     "C3 internal clearance for high-temperature operation.", 312.40, 0.08),
    ("MTR-559", "5HP TEFC Motor 1800RPM",
     "Totally enclosed fan-cooled induction motor, 5 HP at 1800 RPM. "
     "NEMA 184T frame with cast iron housing and Class F insulation. "
     "Rated for 230/460V three-phase operation with a 1.15 service factor.", 875.00, 0.06),
    ("FLT-8801", "Hydraulic Return Filter Assembly",
     "Tank-mounted return line filter with 10 micron cellulose element. "
     "Flow capacity of 40 GPM with integral bypass valve set at 25 PSI. "
     "Visual clogging indicator standard on the die-cast aluminum head.", 156.25, 0.03),
    ("GSK-2210", "Spiral Wound Gasket Set 4\"",
     "Set of six 4-inch spiral wound gaskets, 316SS windings with graphite filler. "
     "Suitable for Class 150 raised-face flanges in steam and hydrocarbon service. "
     "Color coded per ASME B16.20.", 94.80, 0.05),
    ("CPL-4415", "Jaw Coupling 28mm Complete",
     "Three-piece jaw coupling with two sintered iron hubs and NBR spider insert. "
     "Accommodates 28 mm bores with standard keyway. "
     "Torque rating 190 Nm with vibration-damping elastomer element.", 67.90, 0.04),
    ("SEN-9902", "Pressure Transmitter 0-300PSI",
     "Industrial pressure transmitter with 0-300 PSI gauge range and 4-20 mA output. "
     "316L stainless wetted parts with 1/4\" NPT process connection. "
     "Accuracy of 0.25% full scale, CE and UL rated. "
     "Operating temperature -40 to 105 C.", 238.60, 0.05),
    ("HYD-6621", "Hydraulic Cylinder 3\" Bore x 12\" Stroke",
     "Welded cross-tube hydraulic cylinder with 3 inch bore and 12 inch stroke. "
     "3000 PSI working pressure with chrome-plated induction-hardened rod. "
     "SAE ports and greaseable pivot pins on both ends.", 445.00, 0.06),
    ("CMP-1187", "Rotary Screw Air Compressor 10HP",
     "Belt-driven rotary screw compressor delivering 38 CFM at 125 PSI. "
     "10 HP TEFC motor with wye-delta starter mounted on an 80 gallon receiver. "
     "Integrated refrigerated dryer and digital sequencing controller.", 6890.00, 0.05),
    ("ACT-7350", "Pneumatic Rotary Actuator",
     "Rack-and-pinion pneumatic actuator, double-acting, 90 degree rotation. "
     "Output torque 350 in-lb at 80 PSI supply with NAMUR accessory mounting. "
     "Anodized aluminum body with stainless fasteners.", 198.75, 0.04),
    ("HSE-4432", "Hydraulic Hose Assembly 1/2\" x 36\"",
     "Two-wire braided hydraulic hose assembly, 1/2 inch ID by 36 inches long. "
     "Rated 4000 PSI with JIC-8 female swivel ends crimped both sides. "
     "Meets SAE 100R2AT specification.", 38.20, 0.07),
    ("REG-5518", "Air Pressure Regulator 1/2\" NPT",
     "Relieving-type air pressure regulator with 1/2 inch NPT ports. "
     "Adjustable 5-125 PSI outlet range with 220 SCFM flow capacity. "
     "Polycarbonate bowl guard and panel-mount nut included.", 54.35, 0.03),
]

# Popular parts appear in many quotes across years (price history matters).
POPULAR = ["VLV-2043", "PMP-118", "BRG-7742", "SEN-9902"]

CUSTOMERS = [
    "Apex Fabrication LLC", "Great Lakes Processing Inc.", "Thornton Mills",
    "Cascade Water Authority", "Redline Automotive Group", "Pinnacle Foods Plant 7",
    "Harbor Marine Services", "Blue Ridge Paper Co.", "Delta Packaging Systems",
    "Ironwood Mining Corp.",
]

PART_LABELS = ["Part No.", "PN", "Part #", "Item Number", "Part Number", "Item #"]
QTY_LABELS = ["Qty", "Quantity", "Qty.", "Order Qty"]
PRICE_LABELS = ["Unit Price", "Price", "Unit Cost", "Price Each", "Price/Unit"]
DESC_LABELS = ["Description", "Product Description", "Details", "Specification"]


def fmt_price(p, style):
    """Return the messy display string for a price."""
    if style == 0:
        return f"${p:,.2f}"
    if style == 1:
        return f"{p:g} USD"                # 1204.5 USD
    if style == 2:
        return f"USD {p:,.2f}"
    if style == 3:
        return f"${p:.2f} /ea"             # no thousands separator
    return f"{p:,.2f}"                     # bare 1,204.50


def fmt_date(d, style):
    if style == 0:
        return d.strftime("%B %d, %Y")     # March 14, 2023
    if style == 1:
        return d.strftime("%m/%d/%Y")
    if style == 2:
        return d.isoformat()
    return d.strftime("%d-%b-%Y")          # 14-Mar-2023


def quote_number(d, seq, style):
    if style == 0:
        return f"Q-{d.year}-{seq:04d}"
    if style == 1:
        return f"MIS-{d.year % 100}-{seq:04d}"
    return f"QT{d.year % 100}{seq:04d}"


def price_in_year(base, drift, d, rng):
    """Base 2022 price drifted per elapsed year, with small noise, 2dp."""
    years = (d - date(2022, 1, 1)).days / 365.25
    p = base * ((1 + drift) ** years) * rng.uniform(0.99, 1.01)
    return round(p, 2)


def build_quotes(count, rng):
    """Return list of quote dicts (the ground truth)."""
    catalog = {c[0]: c for c in CATALOG}
    start, end = date(2022, 1, 10), date(2026, 6, 30)
    span = (end - start).days
    # Spread dates over the range, keep them sorted so seq numbers make sense.
    days = sorted(rng.sample(range(span), count))
    quotes = []
    for i, dd in enumerate(days):
        d = start + timedelta(days=dd)
        n_items = rng.choices([1, 2, 3, 4], weights=[25, 35, 25, 15])[0]
        # Bias toward popular parts so they recur across years.
        parts = []
        while len(parts) < n_items:
            pool = POPULAR if rng.random() < 0.45 else [c[0] for c in CATALOG]
            p = rng.choice(pool)
            if p not in parts:
                parts.append(p)
        items = []
        for pn in parts:
            _, name, desc, base, drift = catalog[pn]
            price = price_in_year(base, drift, d, rng)
            qty = rng.choices([1, 2, 3, 4, 5, 6, 10, 12, 25],
                              weights=[20, 15, 12, 10, 10, 8, 10, 8, 7])[0]
            # Deliberate missing fields (~12% no description, ~8% no qty)
            has_desc = rng.random() > 0.12
            has_qty = rng.random() > 0.08
            items.append({
                "part_number": pn,
                "product_name": name,
                "description": desc if has_desc else None,
                "unit_price": price,
                "price_style": rng.randrange(5),
                "quantity": qty if has_qty else None,
            })
        style = i % 3  # 0=table, 1=paragraph, 2=letter — even spread
        q = {
            "quote_number": quote_number(d, 100 + i * 7 + rng.randrange(5), rng.randrange(3)),
            "quote_date": d.isoformat(),
            "date_style": rng.randrange(4),
            "customer": rng.choice(CUSTOMERS),
            "layout": ["table", "paragraph", "letter"][style],
            "items": items,
        }
        quotes.append(q)
    # One quote deliberately missing its customer name.
    if count >= 5:
        quotes[rng.randrange(count)]["customer"] = None
    return quotes


# ---------------------------------------------------------- docx renderers ---

def docx_header(doc, q, rng):
    h = doc.add_paragraph()
    r = h.add_run("MERIDIAN INDUSTRIAL SUPPLY CO.")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph("4410 Commerce Park Drive, Toledo, OH 43615  |  (419) 555-0142")
    doc.add_paragraph()
    label = rng.choice(["Quotation", "QUOTE", "Price Quotation", "Sales Quote"])
    p = doc.add_paragraph()
    p.add_run(f"{label}  {q['quote_number']}").bold = True
    doc.add_paragraph(f"Date: {fmt_date(date.fromisoformat(q['quote_date']), q['date_style'])}")
    if q["customer"]:
        doc.add_paragraph(f"{rng.choice(['Customer', 'Prepared for', 'Bill To', 'Client'])}: {q['customer']}")
    doc.add_paragraph()


def render_docx_table(q, path, rng):
    doc = Document()
    docx_header(doc, q, rng)
    pl, ql, prl, dl = (rng.choice(PART_LABELS), rng.choice(QTY_LABELS),
                       rng.choice(PRICE_LABELS), rng.choice(DESC_LABELS))
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for j, txt in enumerate([pl, "Product", dl, ql, prl]):
        hdr[j].paragraphs[0].add_run(txt).bold = True
    for it in q["items"]:
        row = t.add_row().cells
        row[0].text = it["part_number"]
        row[1].text = it["product_name"]
        row[2].text = it["description"] or ""
        row[3].text = "" if it["quantity"] is None else str(it["quantity"])
        row[4].text = fmt_price(it["unit_price"], it["price_style"])
    doc.add_paragraph()
    doc.add_paragraph("Prices valid 30 days. FOB Toledo, OH. Terms Net 30.")
    doc.save(path)


def render_docx_paragraph(q, path, rng):
    doc = Document()
    docx_header(doc, q, rng)
    pl = rng.choice(PART_LABELS)
    prl = rng.choice(PRICE_LABELS)
    ql = rng.choice(QTY_LABELS)
    for i, it in enumerate(q["items"], 1):
        p = doc.add_paragraph()
        p.add_run(f"Line {i}:  {it['product_name']}").bold = True
        doc.add_paragraph(f"{pl}: {it['part_number']}")
        if it["description"]:
            doc.add_paragraph(it["description"])
        bits = []
        if it["quantity"] is not None:
            bits.append(f"{ql}: {it['quantity']}")
        bits.append(f"{prl}: {fmt_price(it['unit_price'], it['price_style'])}")
        doc.add_paragraph("   ".join(bits))
        doc.add_paragraph()
    doc.add_paragraph("This quotation is subject to our standard terms and conditions.")
    doc.save(path)


def render_docx_letter(q, path, rng):
    doc = Document()
    docx_header(doc, q, rng)
    who = q["customer"] or "Valued Customer"
    doc.add_paragraph(f"Dear {who},")
    doc.add_paragraph(
        "Thank you for your inquiry. We are pleased to quote the following "
        "items for your consideration:")
    pl = rng.choice(PART_LABELS)
    for it in q["items"]:
        qty_txt = (f", quantity {it['quantity']}" if it["quantity"] is not None else "")
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{it['product_name']} ({pl} {it['part_number']}){qty_txt} — "
                  f"{fmt_price(it['unit_price'], it['price_style'])} per unit.")
        if it["description"]:
            doc.add_paragraph(it["description"])
    doc.add_paragraph(
        "Pricing is firm for 30 days from the date above. Please reference the "
        "quotation number on any purchase order.")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("J. Whitaker\nSales Engineering\nMeridian Industrial Supply Co.")
    doc.save(path)


# ----------------------------------------------------------- pdf renderers ---

def pdf_styles():
    ss = getSampleStyleSheet()
    ss.add(ParagraphStyle("Co", parent=ss["Title"], fontSize=13, spaceAfter=2))
    ss.add(ParagraphStyle("Small", parent=ss["Normal"], fontSize=8, textColor=colors.grey))
    return ss


def pdf_header(story, ss, q, rng):
    story.append(Paragraph("MERIDIAN INDUSTRIAL SUPPLY CO.", ss["Co"]))
    story.append(Paragraph("4410 Commerce Park Drive, Toledo, OH 43615 | (419) 555-0142", ss["Small"]))
    story.append(Spacer(1, 14))
    label = rng.choice(["Quotation", "QUOTE", "Price Quotation", "Sales Quote"])
    story.append(Paragraph(f"<b>{label}  {q['quote_number']}</b>", ss["Normal"]))
    story.append(Paragraph(
        f"Date: {fmt_date(date.fromisoformat(q['quote_date']), q['date_style'])}", ss["Normal"]))
    if q["customer"]:
        story.append(Paragraph(
            f"{rng.choice(['Customer', 'Prepared for', 'Bill To', 'Client'])}: {q['customer']}",
            ss["Normal"]))
    story.append(Spacer(1, 12))


def render_pdf_table(q, path, rng):
    ss = pdf_styles()
    story = []
    pdf_header(story, ss, q, rng)
    pl, ql, prl, dl = (rng.choice(PART_LABELS), rng.choice(QTY_LABELS),
                       rng.choice(PRICE_LABELS), rng.choice(DESC_LABELS))
    cell = ParagraphStyle("Cell", parent=ss["Normal"], fontSize=8)
    data = [[pl, "Product", dl, ql, prl]]
    for it in q["items"]:
        data.append([
            it["part_number"],
            Paragraph(it["product_name"], cell),
            Paragraph(it["description"] or "", cell),
            "" if it["quantity"] is None else str(it["quantity"]),
            fmt_price(it["unit_price"], it["price_style"]),
        ])
    t = Table(data, colWidths=[0.9 * inch, 1.4 * inch, 3.1 * inch, 0.5 * inch, 1.0 * inch])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))
    story.append(Paragraph("Prices valid 30 days. FOB Toledo, OH. Terms Net 30.", ss["Small"]))
    SimpleDocTemplate(str(path), pagesize=LETTER).build(story)


def render_pdf_paragraph(q, path, rng):
    ss = pdf_styles()
    story = []
    pdf_header(story, ss, q, rng)
    pl, prl, ql = rng.choice(PART_LABELS), rng.choice(PRICE_LABELS), rng.choice(QTY_LABELS)
    for i, it in enumerate(q["items"], 1):
        story.append(Paragraph(f"<b>Line {i}:  {it['product_name']}</b>", ss["Normal"]))
        story.append(Paragraph(f"{pl}: {it['part_number']}", ss["Normal"]))
        if it["description"]:
            story.append(Paragraph(it["description"], ss["Normal"]))
        bits = []
        if it["quantity"] is not None:
            bits.append(f"{ql}: {it['quantity']}")
        bits.append(f"{prl}: {fmt_price(it['unit_price'], it['price_style'])}")
        story.append(Paragraph("   ".join(bits), ss["Normal"]))
        story.append(Spacer(1, 10))
    story.append(Paragraph("This quotation is subject to our standard terms and conditions.", ss["Small"]))
    SimpleDocTemplate(str(path), pagesize=LETTER).build(story)


def render_pdf_letter(q, path, rng):
    ss = pdf_styles()
    story = []
    pdf_header(story, ss, q, rng)
    who = q["customer"] or "Valued Customer"
    story.append(Paragraph(f"Dear {who},", ss["Normal"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "Thank you for your inquiry. We are pleased to quote the following items "
        "for your consideration:", ss["Normal"]))
    story.append(Spacer(1, 6))
    pl = rng.choice(PART_LABELS)
    for it in q["items"]:
        qty_txt = (f", quantity {it['quantity']}" if it["quantity"] is not None else "")
        story.append(Paragraph(
            f"• {it['product_name']} ({pl} {it['part_number']}){qty_txt} — "
            f"{fmt_price(it['unit_price'], it['price_style'])} per unit.",
            ss["Normal"]))
        if it["description"]:
            story.append(Paragraph(it["description"], ss["Normal"]))
        story.append(Spacer(1, 6))
    story.append(Paragraph(
        "Pricing is firm for 30 days from the date above. Please reference the "
        "quotation number on any purchase order.", ss["Normal"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Sincerely,<br/>J. Whitaker<br/>Sales Engineering<br/>"
                           "Meridian Industrial Supply Co.", ss["Normal"]))
    SimpleDocTemplate(str(path), pagesize=LETTER).build(story)


DOCX_RENDER = {"table": render_docx_table, "paragraph": render_docx_paragraph,
               "letter": render_docx_letter}
PDF_RENDER = {"table": render_pdf_table, "paragraph": render_pdf_paragraph,
              "letter": render_pdf_letter}


def main():
    ap = argparse.ArgumentParser(description="Generate fake quote archive")
    ap.add_argument("--count", type=int, default=25, help="number of quotes (default 25)")
    ap.add_argument("--seed", type=int, default=42, help="RNG seed for reproducibility")
    args = ap.parse_args()

    rng = random.Random(args.seed)
    archive = ROOT / "archive"
    data = ROOT / "data"
    archive.mkdir(exist_ok=True)
    data.mkdir(exist_ok=True)

    quotes = build_quotes(args.count, rng)
    key = []
    for q in quotes:
        stem = q["quote_number"].replace("/", "-")
        docx_path = archive / f"{stem}.docx"
        pdf_path = archive / f"{stem}.pdf"
        # Same rng stream would desync docx vs pdf label choices; use a child
        # seed per quote so BOTH renditions pick identical labels/phrasing.
        qseed = rng.randrange(1 << 30)
        DOCX_RENDER[q["layout"]](q, docx_path, random.Random(qseed))
        PDF_RENDER[q["layout"]](q, pdf_path, random.Random(qseed))
        key.append({
            "quote_number": q["quote_number"],
            "quote_date": q["quote_date"],
            "customer": q["customer"],
            "layout": q["layout"],
            "files": [docx_path.name, pdf_path.name],
            "items": [{
                "part_number": it["part_number"],
                "product_name": it["product_name"],
                "description": it["description"],
                "unit_price": it["unit_price"],
                "quantity": it["quantity"],
            } for it in q["items"]],
        })

    with open(data / "answer_key.json", "w", encoding="utf-8") as f:
        json.dump({"quotes": key}, f, indent=2)

    n_items = sum(len(q["items"]) for q in quotes)
    print(f"Generated {len(quotes)} quotes ({n_items} line items) -> {archive}")
    print(f"  {len(quotes)} .docx + {len(quotes)} .pdf files")
    print(f"Answer key -> {data / 'answer_key.json'}  (extractor must never read this)")


if __name__ == "__main__":
    main()
