"""Build the clean professional quote template (template/quote_template.docx).

The template uses {{PLACEHOLDER}} tokens that make_quote.py fills in, plus a
line-items table whose single body row is the row template.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_PATH = Path(__file__).parent / "template" / "quote_template.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # navy


def build_template(path: Path = TEMPLATE_PATH) -> Path:
    path.parent.mkdir(exist_ok=True)
    doc = Document()

    for name, size in (("Normal", 10),):
        doc.styles[name].font.size = Pt(size)
        doc.styles[name].font.name = "Calibri"

    # Letterhead (the seller)
    p = doc.add_paragraph()
    r = p.add_run("MERIDIAN INDUSTRIAL SUPPLY CO.")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT
    doc.add_paragraph("4410 Commerce Park Drive, Toledo, OH 43615  |  (419) 555-0142")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("QUOTATION")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT

    # Quote meta block
    meta = doc.add_table(rows=3, cols=2)
    meta.columns[0].width = Inches(1.5)
    meta.columns[1].width = Inches(5.0)
    for i, (label, token) in enumerate([
            ("Quote No.:", "{{QUOTE_NUMBER}}"),
            ("Date:", "{{DATE}}"),
            ("Customer:", "{{CUSTOMER}}, {{COMPANY}}")]):
        meta.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
        meta.rows[i].cells[1].text = token
    doc.add_paragraph()

    # Line items table: header + ONE body row of tokens (cloned per item)
    t = doc.add_table(rows=2, cols=6)
    t.style = "Table Grid"
    headers = ["Part Number", "Product", "Description", "Qty", "Unit Price", "Ext. Price"]
    for j, h in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = ACCENT
    tokens = ["{{PART}}", "{{NAME}}", "{{DESC}}", "{{QTY}}", "{{PRICE}}", "{{EXT}}"]
    for j, tok in enumerate(tokens):
        t.rows[1].cells[j].text = tok
    widths = [1.0, 1.3, 2.6, 0.5, 0.9, 0.9]
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Inches(w)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Total: {{TOTAL}}").bold = True

    doc.add_paragraph()
    doc.add_paragraph("Prices are valid for 30 days from the date above. "
                      "FOB Toledo, OH. Payment terms Net 30.")
    doc.add_paragraph("Thank you for the opportunity to quote your requirements.")

    doc.save(str(path))
    return path


if __name__ == "__main__":
    print(f"Template written to {build_template()}")
