"""Generate a new quote from the indexed archive.

Usage:
    python make_quote.py --parts VLV-2043,PMP-118 --customer "Jane Doe" --company "Apex Fabrication LLC"

For each requested part number the MOST RECENT archived occurrence wins
(tie-break rule from rules.md). The result is written to output/ using
template/quote_template.docx (auto-created if missing).

If a part is not found, the script fails gracefully: it lists what WAS
found, suggests close matches (fuzzy), and writes nothing unless
--allow-partial is given.
"""

import argparse
import copy
import difflib
import sys
from datetime import date
from pathlib import Path

from docx import Document

from common import ROOT, DB_PATH, db_connect
from template import TEMPLATE_PATH, build_template

OUTPUT_DIR = ROOT / "output"


def lookup_part(con, part):
    """Most recent line item for a part number (case-insensitive).

    The archive stores each quote as both .docx and .pdf, so the same quote
    can match twice — ordering by date then id and taking the first row
    dedupes that naturally.
    """
    row = con.execute(
        """SELECT li.part_number, li.product_name, li.description,
                  li.unit_price, q.quote_date, q.quote_number, q.file
           FROM line_items li JOIN quotes q ON q.id = li.quote_id
           WHERE UPPER(li.part_number) = UPPER(?) AND q.quote_date IS NOT NULL
           ORDER BY q.quote_date DESC, q.id ASC
           LIMIT 1""", (part.strip(),)).fetchone()
    if row is None:
        return None
    keys = ("part_number", "product_name", "description", "unit_price",
            "quote_date", "quote_number", "source_file")
    hit = dict(zip(keys, row))
    # Price always comes from the most recent quote, but if that quote was
    # missing the name/description, backfill from the most recent occurrence
    # that has one (archives are messy; the product itself hasn't changed).
    for field in ("product_name", "description"):
        if hit[field] is None:
            fb = con.execute(
                f"""SELECT li.{field}
                    FROM line_items li JOIN quotes q ON q.id = li.quote_id
                    WHERE UPPER(li.part_number) = UPPER(?)
                      AND li.{field} IS NOT NULL AND q.quote_date IS NOT NULL
                    ORDER BY q.quote_date DESC, q.id ASC LIMIT 1""",
                (part.strip(),)).fetchone()
            if fb:
                hit[field] = fb[0]
    return hit


def suggest(con, part, n=3):
    known = [r[0] for r in con.execute(
        "SELECT DISTINCT part_number FROM line_items WHERE part_number IS NOT NULL")]
    return difflib.get_close_matches(part.upper(), known, n=n, cutoff=0.55)


def next_quote_number(out_dir: Path) -> str:
    year = date.today().year
    existing = len(list(out_dir.glob(f"Q-{year}-*.docx")))
    return f"Q-{year}-{existing + 1:04d}"


def replace_tokens(doc, mapping):
    """Replace {{TOKEN}} strings in all paragraphs and table cells."""
    def fix(paragraph):
        text = paragraph.text
        if "{{" not in text:
            return
        for k, v in mapping.items():
            text = text.replace(k, v)
        # Collapse runs: keep first run's formatting.
        for run in paragraph.runs[1:]:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)

    for p in doc.paragraphs:
        fix(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix(p)


def fill_items_table(doc, items, quantities):
    """Clone the token row of the line-items table once per item."""
    table = next(t for t in doc.tables
                 if "{{PART}}" in t.rows[1].cells[0].text)
    template_row = table.rows[1]
    total = 0.0
    for item, qty in zip(items, quantities):
        new_tr = copy.deepcopy(template_row._tr)
        template_row._tr.addprevious(new_tr)
        row = table.rows[len(table.rows) - 2]  # the row just inserted
        ext = (item["unit_price"] or 0.0) * qty
        total += ext
        values = [item["part_number"], item["product_name"] or "",
                  item["description"] or "", str(qty),
                  f"${item['unit_price']:,.2f}" if item["unit_price"] is not None else "—",
                  f"${ext:,.2f}"]
        for cell, val in zip(row.cells, values):
            # Clear existing text, set value with the template row's formatting.
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ""
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = val
            else:
                cell.paragraphs[0].add_run(val)
    # Remove the token row.
    template_row._tr.getparent().remove(template_row._tr)
    return total


def generate(parts, customer, company, quantities=None, db_path=DB_PATH,
             out_dir=OUTPUT_DIR, allow_partial=False, today=None):
    """Core generation logic (importable for tests). Returns (path, found, missing)."""
    con = db_connect(db_path)
    quantities = quantities or [1] * len(parts)
    found, missing = [], []
    for part in parts:
        hit = lookup_part(con, part)
        if hit:
            found.append(hit)
        else:
            missing.append((part, suggest(con, part)))

    if missing and not (allow_partial and found):
        return None, found, missing
    if not found:
        return None, found, missing

    out_dir.mkdir(exist_ok=True)
    if not TEMPLATE_PATH.exists():
        build_template()
    doc = Document(str(TEMPLATE_PATH))

    qnum = next_quote_number(out_dir)
    today = today or date.today()
    qty_used = [q for p, q in zip(parts, quantities)
                if any(f["part_number"].upper() == p.strip().upper() for f in found)]
    total = fill_items_table(doc, found, qty_used)
    replace_tokens(doc, {
        "{{COMPANY}}": company,
        "{{CUSTOMER}}": customer,
        "{{DATE}}": today.strftime("%B %d, %Y"),
        "{{QUOTE_NUMBER}}": qnum,
        "{{TOTAL}}": f"${total:,.2f}",
    })
    out_path = out_dir / f"{qnum}.docx"
    doc.save(str(out_path))
    con.close()
    return out_path, found, missing


def main():
    ap = argparse.ArgumentParser(description="Generate a new quote from the archive")
    ap.add_argument("--parts", required=True,
                    help="comma-separated part numbers, e.g. VLV-2043,PMP-118")
    ap.add_argument("--customer", required=True, help="customer contact name")
    ap.add_argument("--company", required=True, help="customer company name")
    ap.add_argument("--qty", default=None,
                    help="comma-separated quantities matching --parts (default 1 each)")
    ap.add_argument("--allow-partial", action="store_true",
                    help="still generate the quote if some parts are missing")
    args = ap.parse_args()

    if not DB_PATH.exists():
        raise SystemExit(f"No database at {DB_PATH}. Run index_archive.py first.")

    parts = [p.strip() for p in args.parts.split(",") if p.strip()]
    quantities = ([int(q) for q in args.qty.split(",")] if args.qty
                  else [1] * len(parts))
    if len(quantities) != len(parts):
        raise SystemExit("--qty must have one value per part in --parts")

    out_path, found, missing = generate(
        parts, args.customer, args.company, quantities,
        allow_partial=args.allow_partial)

    if found:
        print("Found in archive (most recent occurrence wins):")
        for f in found:
            print(f"  {f['part_number']:<10} {f['product_name']:<40} "
                  f"${f['unit_price']:,.2f}  (from {f['quote_number']}, {f['quote_date']})")
    if missing:
        print("\nNOT FOUND:")
        for part, sugg in missing:
            hint = f"  — did you mean: {', '.join(sugg)}?" if sugg else ""
            print(f"  {part}{hint}")

    if out_path:
        print(f"\nQuote written to {out_path}")
    elif missing:
        print("\nNo quote generated. Fix the part numbers above, or re-run "
              "with --allow-partial to generate with the found items only.")
        sys.exit(1)
    else:
        print("\nNo quote generated: none of the requested parts were found.")
        sys.exit(1)


if __name__ == "__main__":
    main()
